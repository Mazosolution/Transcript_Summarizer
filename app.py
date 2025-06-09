import streamlit as st
import google.generativeai as genai
import os
from dotenv import load_dotenv
import io
import base64
from docx import Document
import re

# Load environment variables from .env file
load_dotenv()

# Configure Google Gemini API
def configure_gemini_api():
    """Configure the Google Gemini API with the API key from .env."""
    try:
        api_key = os.getenv("GOOGLE_API_KEY")
        if not api_key:
            st.error("Google API key not found in .env file.")
            st.stop()
        
        genai.configure(api_key=api_key)
    except Exception as e:
        st.error(f"Error configuring API: {str(e)}")
        st.stop()

# Function to read different file types
def read_file_content(uploaded_file):
    """Read content from different file types."""
    try:
        if uploaded_file.type == 'application/pdf':
            try:
                import PyPDF2
            except ImportError:
                st.error("PyPDF2 is required to read PDF files.")
                return None
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            content = ' '.join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
        elif uploaded_file.type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            doc = Document(uploaded_file)
            content = ' '.join([para.text for para in doc.paragraphs])
        elif uploaded_file.type == 'text/plain':
            content = uploaded_file.getvalue().decode('utf-8')
        else:
            st.warning(f"Unsupported file type: {uploaded_file.type}. Attempting to read as text.")
            try:
                content = uploaded_file.getvalue().decode('utf-8')
            except Exception:
                st.error(f"Could not read file '{uploaded_file.name}' as text.")
                return None
        return content
    except Exception as e:
        st.error(f"Error reading file '{uploaded_file.name}': {str(e)}")
        return None

# Function to generate comprehensive interview analysis
def generate_interview_analysis(transcript_text, job_description, decision_levels):
    """Generate a structured analysis of the interview transcript using Gemini."""
    try:
        # Update to use the recommended model
        model = genai.GenerativeModel('gemini-1.5-flash')
    except Exception as e:
        st.error(f"Error creating Generative Model: {str(e)}.")
        return {"analysis": "Could not configure the AI model."}
    
    job_context = f"""
Job Description:
---
{job_description}
---
Evaluate the candidate's responses, skills, and experience discussed in the transcript against this job description. Apply objective assessment criteria when evaluating candidates.
"""

    # Add rating guidelines to ensure proper decision making
    rating_guidelines = """
Rating to Decision Level Guidelines:
* 5/5: Excellent candidate fit - SELECT
* 4/5: Strong candidate with minor gaps - SELECT
* 3/5: Acceptable candidate with some gaps - HOLD
* 2/5: Candidate with significant gaps - HOLD (if gaps are in non-critical areas) or REJECT (if gaps are in critical requirements)
* 1/5: Poor candidate fit - REJECT

When making your final decision recommendation, please ensure it aligns with the overall rating.
Please ensure a balanced assessment giving candidates fair consideration.
"""

    prompt = f"""Analyze the following interview transcript, comparing it against the provided job description. Provide a comprehensive, structured summary with the following sections:

1.  **Interview Overview**
    *   Candidate Name (if mentioned, otherwise "Not Mentioned")
    *   Position/Role Applied For (based on JD or transcript)
    *   Interview Type (e.g., Technical Screen, Behavioral, Final Round - infer if possible)
    *   Date/Time (if mentioned, otherwise "Not Mentioned")

2.  **Candidate Background Summary**
    *   Brief summary of claimed relevant experience (years, key areas).
    *   Mentioned skills relevant to the Job Description.

3.  **Key Questions and Candidate Responses**
    *   Identify and list the main technical or behavioral questions asked by the interviewer(s).
    *   For each key question, summarize the candidate's response.
    *   Analyze the quality and depth of the candidate's answers objectively.
    *   Count and mention the number of distinct technical/coding questions asked.
    *   Count and mention the number of questions the candidate provided good or adequate answers to.

4.  **Job Description Alignment Analysis**
    *   Compare the candidate's stated experience and demonstrated skills against key requirements in the Job Description.
    *   Objectively evaluate where the candidate's skills match or don't match the requirements.

5.  **Communication and Professionalism**
    *   Assess the candidate's communication style (clarity, conciseness).
    *   Note their apparent confidence level during the interview.

6.  **Interviewer Performance (Brief)**
    *   Assess if the interviewer's questions effectively covered key areas mentioned in the job description.
    *   Note the structure and flow of the interview.
    *   Calculate the approximate number of questions asked by the interviewer.

7.  **Overall Assessment**
    *   **Overall Rating (Score: X/5):** Provide a numerical rating from 1 (Poor Fit) to 5 (Excellent Fit) based on objective assessment.
    *   **Justification:** Explain the rating with factual observations from the interview.

8.  **Final Decision Recommendation**
    *   **Recommendation:** Choose ONE level: {', '.join(decision_levels.keys())}.
        *   SELECT: Candidate who objectively meets key requirements for the role.
        *   HOLD: Candidate needs further evaluation in specific areas.
        *   REJECT: Candidate does not meet essential requirements for this role.
    *   **Supporting Points:** Provide 3-5 concise bullet points supporting the recommendation.
    *   **Confidence Level:** State confidence in the recommendation (Low, Medium, High).

{rating_guidelines}

Transcript:
---
{transcript_text}
---

{job_context}

**Instructions for AI:**
*   Base your evaluation on objective criteria from the job description.
*   Give equal weight to strengths and areas for improvement.
*   Assess the candidate based on demonstrated skills and experience.
*   Use the full range of ratings (1-5) as appropriate.
*   Structure the output clearly using the headings provided above. Use Markdown for formatting.
*   Ensure your final recommendation (SELECT/HOLD/REJECT) aligns with your overall rating.
*   Do not be overly critical. Consider the candidate's potential to grow into the role.
"""

    try:
        response = model.generate_content(prompt)
        analysis_text = ""
        if hasattr(response, 'text'):
            analysis_text = response.text
        elif hasattr(response, 'parts') and response.parts:
            analysis_text = "".join(part.text for part in response.parts if hasattr(part, 'text'))
        else:
            analysis_text = "Error: Could not parse AI response."
        
        # Add a review step to check for bias
        review_prompt = f"""
Review the following interview analysis and check if the decision recommendation appears biased or unreasonably strict:

Analysis:
{analysis_text}

If the decision seems overly harsh given the candidate's strengths and the overall rating, please provide a more balanced reassessment.
Focus specifically on whether the final decision (SELECT/HOLD/REJECT) aligns with the numerical rating (1-5).

Guidelines:
- Rating 4-5/5 should typically correspond to SELECT
- Rating 3/5 should typically correspond to HOLD
- Rating 1-2/5 should typically correspond to REJECT, but consider HOLD for 2/5 if gaps are not in core requirements

Provide only a brief correction if needed, or confirm the original assessment if it seems fair.
"""
        
        try:
            review_response = model.generate_content(review_prompt)
            review_text = ""
            if hasattr(review_response, 'text'):
                review_text = review_response.text
            elif hasattr(review_response, 'parts') and review_response.parts:
                review_text = "".join(part.text for part in review_response.parts if hasattr(part, 'text'))
            
            # Only include the review if it suggests a correction
            if "correction" in review_text.lower() or "reassess" in review_text.lower() or "reconsider" in review_text.lower():
                analysis_text += "\n\n## Bias Check\n" + review_text
        except Exception:
            # If review fails, continue with original analysis
            pass
             
        return {
            "analysis": analysis_text
        }
    except Exception as e:
        st.error(f"Error generating analysis: {str(e)}")
        return {"analysis": f"Could not generate analysis. Error: {str(e)}"}

# Function to create a downloadable Word document
def create_word_doc(analysis_text):
    """Create a Word document from the analysis text and return a download link."""
    try:
        doc = Document()
        doc.add_heading('Interview Analysis Report', 0)
        lines = analysis_text.split('\n')
        current_heading_level = 0
        
        for line in lines:
            stripped_line = line.strip()
            if not stripped_line:
                continue

            if stripped_line.startswith('**') and stripped_line.endswith('**'):
                heading_text = stripped_line.strip('*').strip()
                if re.match(r'^\d+\.\s+', heading_text):
                    doc.add_heading(heading_text, level=1)
                    current_heading_level = 1
                elif heading_text in ["Overall Rating", "Justification", "Recommendation", "Supporting Points", "Confidence Level"]:
                    doc.add_heading(heading_text, level=2)
                    current_heading_level = 2
                else:
                    p = doc.add_paragraph()
                    p.add_run(heading_text).bold = True
                    current_heading_level = 0

            elif re.match(r'^\d+\.\s+', stripped_line):
                doc.add_heading(stripped_line, level=1)
                current_heading_level = 1
            
            elif stripped_line.startswith(('* ', '- ', '+ ')):
                text = stripped_line[2:]
                doc.add_paragraph(text, style='List Bullet')
                current_heading_level = 0

            else:
                doc.add_paragraph(stripped_line)
                current_heading_level = 0

        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        b64_doc = base64.b64encode(doc_io.read()).decode('utf-8')
        
        return f'data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64_doc}'
    except Exception as e:
        st.error(f"Error creating Word document: {str(e)}")
        return None

# Configure decision levels with colors and descriptions
def get_decision_levels():
    return {
        "SELECT": {"color": "#34A853", "description": "Candidate who meets key requirements for the role."},
        "HOLD": {"color": "#FBBC04", "description": "Candidate needs further evaluation in specific areas."},
        "REJECT": {"color": "#EA4335", "description": "Candidate does not meet essential requirements for this role."}
    }

# Extract Overall Rating (1-5) from analysis text
def extract_overall_rating(analysis_text):
    """Extracts the overall rating (1-5) from the analysis text."""
    rating = {"Overall Rating": 0}
    match = re.search(r"Overall Rating(?:.*?Score\s*:\s*)?(\b[1-5](?:\.\d)?\b)\s*/\s*5", analysis_text, re.IGNORECASE | re.DOTALL)
    
    if match:
        try:
            rating_value = float(match.group(1))
            if 1 <= rating_value <= 5:
                rating["Overall Rating"] = rating_value
        except (ValueError, IndexError):
            pass

    if rating["Overall Rating"] == 0:
        if "7. Overall Assessment" in analysis_text:
            assessment_section = analysis_text.split("7. Overall Assessment")[1].split("8. Final Decision Recommendation")[0]
            numbers = re.findall(r'\b([1-5](?:\.\d)?)\b', assessment_section)
            if numbers:
                try:
                    rating["Overall Rating"] = float(numbers[0])
                except ValueError:
                    pass

    return rating

# Function to determine decision level from analysis text
def extract_decision_level(analysis_text, decision_levels):
    """Extracts the final decision level and description from the analysis text with improved accuracy."""
    # Default to HOLD as a middle ground
    result = {
        "level": "HOLD",
        "color": decision_levels["HOLD"]["color"],
        "description": decision_levels["HOLD"]["description"]
    }
    
    # Look for specific recommendation section
    recommendation_patterns = [
        r"(?:Recommendation|RECOMMENDATION):\s*([A-Za-z]+)",
        r"(?:Final Decision|FINAL DECISION)(?:.*?)(?:Recommendation|RECOMMENDATION):\s*([A-Za-z]+)",
        r"8\.\s+(?:.*?)(?:Recommendation|RECOMMENDATION):\s*([A-Za-z]+)"
    ]
    
    for pattern in recommendation_patterns:
        match = re.search(pattern, analysis_text, re.IGNORECASE)
        if match:
            level_found = match.group(1).upper().strip()
            # Check if the found level is in our defined levels
            for key in decision_levels.keys():
                if key in level_found:
                    result = {
                        "level": key,
                        "color": decision_levels[key]["color"],
                        "description": decision_levels[key]["description"]
                    }
                    return result
    
    # If we're still using the default, check for bias correction
    if result["level"] == "HOLD" and "## Bias Check" in analysis_text:
        bias_section = analysis_text.split("## Bias Check")[1]
        for key in decision_levels.keys():
            if key in bias_section.upper():
                result = {
                    "level": key,
                    "color": decision_levels[key]["color"],
                    "description": decision_levels[key]["description"]
                }
                return result
    
    # As a backup, correlate with overall rating
    overall_rating = extract_overall_rating(analysis_text).get("Overall Rating", 0)
    if overall_rating > 0:
        # Map rating to decision level
        if overall_rating >= 4:
            result = {
                "level": "SELECT",
                "color": decision_levels["SELECT"]["color"],
                "description": decision_levels["SELECT"]["description"]
            }
        elif overall_rating >= 2.5:
            result = {
                "level": "HOLD",
                "color": decision_levels["HOLD"]["color"],
                "description": decision_levels["HOLD"]["description"]
            }
        else:
            result = {
                "level": "REJECT",
                "color": decision_levels["REJECT"]["color"],
                "description": decision_levels["REJECT"]["description"]
            }
            
    return result

# Streamlit App
def main():
    st.set_page_config(
        page_title="Interview Transcript Analyzer", 
        page_icon="📋",
        layout="wide"
    )
    
    st.title("🤝 Interview Transcript Analyzer")
    st.markdown("Analyzes interview transcripts against job descriptions, providing objective assessment and recommendations.")

    configure_gemini_api()
    decision_levels = get_decision_levels()
    
    col1, col2 = st.columns(2)
    
    with col1:
        transcript_file = st.file_uploader(
            "Upload Interview Transcript", 
            type=['txt', 'pdf', 'docx'],
            help="Upload a text, PDF, or DOCX file containing the interview transcript."
        )
    
    with col2:
        job_desc_file = st.file_uploader(
            "Upload Job Description (Required)", 
            type=['txt', 'pdf', 'docx'],
            help="Upload a text, PDF, or DOCX file containing the job description."
        )
    
    with st.expander("Decision Level Guide"):
        st.subheader("Understanding Decision Levels")
        for level, info in decision_levels.items():
            st.markdown(f"<div style='padding:8px;margin-bottom:8px;border-radius:4px;background-color:{info['color']}20;border-left:4px solid {info['color']}'><strong style='color:{info['color']}'>{level}</strong>: {info['description']}</div>", unsafe_allow_html=True)
        
        # Add rating to decision level mapping explanation
        st.subheader("Rating to Decision Mapping")
        st.markdown("""
        <div style='padding:10px;border-radius:4px;background-color:#f0f0f0;'>
        <p><strong>How ratings translate to decisions:</strong></p>
        <ul>
        <li><strong>Rating 4-5:</strong> Typically SELECT - Candidate meets requirements</li>
        <li><strong>Rating 3:</strong> Typically HOLD - Candidate has potential but needs further evaluation</li>
        <li><strong>Rating 1-2:</strong> Typically REJECT - Candidate does not meet essential requirements</li>
        </ul>
        </div>
        """, unsafe_allow_html=True)
    
    analyze_button = st.button("Analyze Transcript", type="primary", disabled=(not transcript_file or not job_desc_file))

    if 'analysis_result' not in st.session_state:
        st.session_state.analysis_result = None
    if 'error_message' not in st.session_state:
        st.session_state.error_message = None

    if analyze_button:
        st.session_state.analysis_result = None
        st.session_state.error_message = None

        transcript_text = read_file_content(transcript_file)
        job_description = read_file_content(job_desc_file)
        
        if transcript_text and job_description:
            st.write("🔍 Analyzing Interview Transcript...")
            progress_bar = st.progress(0)
            status_text = st.empty()
            status_text.text("Generating objective analysis...")
            
            with st.spinner("AI is processing the transcript and job description..."):
                analysis_result = generate_interview_analysis(transcript_text, job_description, decision_levels)
                st.session_state.analysis_result = analysis_result
            
            progress_bar.progress(100)
            status_text.text("Analysis Complete!")

        else:
            error_msgs = []
            if not transcript_text:
                error_msgs.append("Could not read the transcript file. Please check the file format or content.")
            if not job_description:
                error_msgs.append("Could not read the job description file. Please check the file format or content.")
            st.session_state.error_message = "\n".join(error_msgs)
            st.error(st.session_state.error_message)

    if st.session_state.analysis_result:
        analysis_content = st.session_state.analysis_result.get("analysis", "")
        
        if "Could not generate analysis" in analysis_content or "Error:" in analysis_content:
            st.error(f"Failed to generate analysis: {analysis_content}")
        elif analysis_content:
            st.header("📄 Interview Analysis Report")
            
            analysis_text = analysis_content
            rating_data = extract_overall_rating(analysis_text)
            overall_rating = rating_data.get("Overall Rating", 0)
            decision = extract_decision_level(analysis_text, decision_levels)
            
            st.markdown("---")
            metrics_col1, metrics_col2 = st.columns([1, 2])
            
            with metrics_col1:
                st.markdown("#### Overall Rating")
                st.metric(label="Overall Rating", value=f"{overall_rating} / 5")
                
                if overall_rating >= 4:
                    rating_color = "#34A853"
                elif overall_rating >= 2.5:
                    rating_color = "#FBBC04"
                else:
                    rating_color = "#EA4335"
                
                st.markdown(f"""
                <div style="width: 90%; background-color: #e0e0e0; border-radius: 5px; margin-top: 5px;">
                    <div style="width: {float(overall_rating) * 20}%; background-color: {rating_color}; height: 20px; border-radius: 5px; text-align: center; color: white; font-weight: bold; line-height: 20px;">
                        {overall_rating}/5
                    </div>
                </div>
                """, unsafe_allow_html=True)

            with metrics_col2:
                st.markdown(f"#### Final Decision Recommendation")
                st.markdown(f"<div style='padding:10px; background-color:{decision['color']}20; border-left: 5px solid {decision['color']}; border-radius: 5px; margin-bottom: 10px;'>"
                            f"<strong style='color:{decision['color']}; font-size: 1.1em;'>{decision['level']}</strong>: {decision['description']}"
                            f"</div>", unsafe_allow_html=True)
                
                # Add a note if decision and rating seem misaligned
                if (overall_rating >= 4 and decision['level'] != "SELECT") or \
                   (overall_rating >= 2.5 and overall_rating < 4 and decision['level'] != "HOLD") or \
                   (overall_rating < 2.5 and decision['level'] != "REJECT"):
                    st.markdown(f"<div style='padding:10px; background-color:#FFF3CD; border-left: 5px solid #FFC107; border-radius: 5px;'>"
                                f"<strong>Note:</strong> The decision recommendation may not align with the overall rating. "
                                f"Consider reviewing the analysis to ensure consistency.</div>", unsafe_allow_html=True)

            st.markdown("---")
            st.markdown("### Detailed Analysis")
            st.markdown(analysis_text) 
            
            st.markdown("---")
            st.markdown(f"### Download Report")
            word_link = create_word_doc(analysis_text)
            if word_link:
                st.download_button(
                    label="Download Analysis as Word Document (.docx)",
                    data=base64.b64decode(word_link.split(',')[1]),
                    file_name="interview_analysis_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.warning("Could not generate Word document for download.")
                 
    elif st.session_state.error_message:
        st.error(st.session_state.error_message)

    elif not analyze_button and not st.session_state.analysis_result and not st.session_state.error_message:
        if not transcript_file:
            st.info("Please upload an interview transcript file to begin.")
        if not job_desc_file:
            st.warning("Job description is required. Please upload a job description file.")
        if transcript_file and job_desc_file:
            st.info("Click the 'Analyze Transcript' button to generate the report.")

if __name__ == "__main__":
    main()
